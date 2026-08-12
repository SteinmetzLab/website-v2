"""Write a refreshed CV from cv20250905.docx with the publication tables brought current.

Each publication is ONE paragraph inside a three-cell row (year | entry | journal). The
entry paragraph is not plain text: it is a sequence of runs --

    [author runs, with "Steinmetz" bold] <w:br/> <w:hyperlink>[title run, Hyperlink style]</>
    optionally followed by "|" and an italic "preprint" / "review"

so new entries have to be assembled run by run. Writing the text straight into the
paragraph's w:t elements, which is the obvious thing to try, silently collapses the bold
author name and the hyperlink into one unstyled run.

Usage:  python tools/update_cv.py [--pdf]
"""
import copy
import datetime as dt
import re
import sys
from pathlib import Path

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
CVDIR = Path(r"D:/Dropbox/cv/cv")
SRC = CVDIR / "cv20250905.docx"
DST = CVDIR / "cv20260811_v2.docx"

# (year, authors, title, journal, url, suffix) -- newest first within each table
SENIOR_NEW = [
    ("2026", "Ye, Ladd, MacKenzie, Kolich, Li, Birman, Bull, Daigle, Tasic, Zeng, Steinmetz",
     "Brain-wide Topographic Coordination of Rotating Waves", "Science",
     "https://doi.org/10.1126/science.adx1369", ""),
    ("", "Shaker, Schroeter, Birman, Steinmetz",
     "The Midbrain Reticular Formation in Contextual Control of Perceptual Decisions", "Neuron",
     "https://doi.org/10.1016/j.neuron.2026.04.003", ""),
    ("", "Siegle, Steinmetz",
     "Large-scale Electrophysiology at Single-spike Resolution", "Nature Rev Neurosci",
     "https://doi.org/10.1038/s41583-026-01042-4", "review"),
    ("", "Roth, Chapuis, Winter, The International Brain Laboratory, \u2026, Horwitz, Steinmetz",
     "A Flexible Quality Metric for Electrophysiological Recordings Across Brain Regions and Species",
     "bioRxiv", "https://doi.org/10.1101/2026.03.06.710130", "preprint"),
    ("2025", "Lu, Li, Ladd, Matveev, Deole, Shea-Brown, Kutz, Steinmetz",
     "Benchmarking Probabilistic Time Series Forecasting Models on Neural Activity", "arXiv",
     "https://doi.org/10.48550/arXiv.2510.18037", "preprint"),
]
MIDDLE_NEW = [
    ("2026", "Hjort, Garrett, Gordon, Ancell, Trzeciak, Lu, Bruchas, Witten, Steinmetz, Stuber",
     "Prefrontal to Ventral Tegmental Area Dynamics Drive Contingency Degradation", "Nature",
     "https://doi.org/10.1038/s41586-026-10443-5", ""),
]

JOURNAL_UPDATES = [
    (0, "Ultra-High Density Neuropixels", "Neuron"),
    (1, "Neuropixels Opto", "Nature Methods"),
    (1, "A Multimodal Approach for Visualization", "Nature Comm"),
    (1, "Active Filtering of Sequences", "Neuron"),
]
# published since the last CV, so the link should no longer point at the preprint
RELINK = [(0, "Ultra-High Density Neuropixels", "https://doi.org/10.1016/j.neuron.2025.08.030")]
RETITLE = [
    (1, "A Multimodal Approach for Visualization",
     "A Multimodal Approach for Visualizing and Identifying Electrophysiological Cell Types",
     "https://doi.org/10.1038/s41467-026-71331-0"),
    (1, "Active Filtering of Sequences",
     "Recurrent Cortical Networks Encode Natural Sensory Statistics via Sequence Filtering",
     "https://doi.org/10.1016/j.neuron.2025.12.024"),
]
DROP_PREPRINT_LABEL = [(1, "Neuropixels Opto"), (1, "A Multimodal Approach for Visualizing"),
                       (1, "Recurrent Cortical Networks Encode")]
DELETE = [(0, "Brain-wide Topographic Coordination of Traveling Spiral Waves")]


def text_of(el):
    return re.sub(r"\s+", " ", "".join(t.text or "" for t in el.iter(NS + "t"))).strip()


class Templates:
    """Run shapes lifted from a real entry, so new runs match the document exactly."""

    def __init__(self, para):
        self.plain = self.bold = self.italic = self.link_run = self.br = self.hyperlink = None
        for r in para.findall(NS + "r"):
            rPr = r.find(NS + "rPr")
            if r.find(NS + "br") is not None and self.br is None:
                self.br = copy.deepcopy(r)
            elif rPr is not None and rPr.find(NS + "b") is not None and self.bold is None:
                self.bold = copy.deepcopy(r)
            elif rPr is not None and rPr.find(NS + "i") is not None and self.italic is None:
                self.italic = copy.deepcopy(r)
            elif self.plain is None and r.find(NS + "t") is not None:
                self.plain = copy.deepcopy(r)
        hl = para.find(NS + "hyperlink")
        if hl is not None:
            self.hyperlink = copy.deepcopy(hl)
            self.link_run = copy.deepcopy(hl.find(NS + "r"))
        if self.italic is None and self.plain is not None:
            self.italic = copy.deepcopy(self.plain)
            rPr = self.italic.find(NS + "rPr")
            rPr.append(rPr.makeelement(NS + "i", {}))

    def run(self, kind, text):
        r = copy.deepcopy(getattr(self, kind))
        for t in r.findall(NS + "t"):
            r.remove(t)
        t = r.makeelement(NS + "t", {})
        t.text = text
        t.set(XML_SPACE, "preserve")
        r.append(t)
        return r


def build_entry(para, tpl, authors, title, url, suffix, part):
    """Replace a paragraph's content with a properly run-structured entry."""
    for child in list(para):
        if child.tag != NS + "pPr":
            para.remove(child)

    head, sep, tail = authors.partition("Steinmetz")   # the CV bolds its own name throughout
    if head:
        para.append(tpl.run("plain", " " + head))
    if sep:
        para.append(tpl.run("bold", sep))
    if tail:
        para.append(tpl.run("plain", tail))
    para.append(copy.deepcopy(tpl.br))

    hl = copy.deepcopy(tpl.hyperlink)
    for r in hl.findall(NS + "r"):
        hl.remove(r)
    hl.append(tpl.run("link_run", title))
    hl.set(R_NS + "id", part.relate_to(url, RT.HYPERLINK, is_external=True))
    para.append(hl)

    if suffix:
        para.append(tpl.run("plain", " | "))
        para.append(tpl.run("italic", suffix))


def find_row(tbl, fragment):
    for i, tr in enumerate(tbl._tbl.findall(NS + "tr")):
        if fragment.lower() in text_of(tr).lower():
            return i
    return None


def entry_paragraph(tr, fragment):
    for tc in tr.findall(NS + "tc"):
        for p in tc.findall(NS + "p"):
            if fragment.lower() in text_of(p).lower():
                return p
    return None


def set_cell_text(cell, value):
    ts = list(cell.iter(NS + "t"))
    ts[0].text = value
    ts[0].set(XML_SPACE, "preserve")
    for extra in ts[1:]:
        extra.text = ""


def main():
    d = docx.Document(str(SRC))
    part = d.part
    tables = {0: d.tables[0], 1: d.tables[1]}
    tpl = Templates(tables[0]._tbl.findall(NS + "tr")[0].findall(NS + "tc")[1].find(NS + "p"))
    assert all([tpl.plain, tpl.bold, tpl.br, tpl.hyperlink]), "template runs not found"
    report = []

    for ti, frag, journal in JOURNAL_UPDATES:
        tbl = tables[ti]
        i = find_row(tbl, frag)
        trs = tbl._tbl.findall(NS + "tr")
        cell = trs[i].findall(NS + "tc")[-1]
        if not text_of(cell):                      # the journal sits on the authors' row
            cell = trs[i - 1].findall(NS + "tc")[-1]
        old = text_of(cell)
        set_cell_text(cell, journal)
        report.append(f"  t{ti}: {frag[:34]!r} journal {old!r} -> {journal!r}")

    for ti, frag, url in RELINK:
        tbl = tables[ti]
        tr = tbl._tbl.findall(NS + "tr")[find_row(tbl, frag)]
        hl = entry_paragraph(tr, frag).find(NS + "hyperlink")
        hl.set(R_NS + "id", part.relate_to(url, RT.HYPERLINK, is_external=True))
        report.append(f"  t{ti}: relinked {frag[:32]!r} to the published version")

    for ti, frag, title, url in RETITLE:
        tbl = tables[ti]
        tr = tbl._tbl.findall(NS + "tr")[find_row(tbl, frag)]
        p = entry_paragraph(tr, frag)
        hl = p.find(NS + "hyperlink")
        set_cell_text(hl, title)
        hl.set(R_NS + "id", part.relate_to(url, RT.HYPERLINK, is_external=True))
        report.append(f"  t{ti}: retitled + relinked -> {title[:46]!r}")

    for ti, frag in DROP_PREPRINT_LABEL:
        tbl = tables[ti]
        i = find_row(tbl, frag)
        if i is None:
            continue
        p = entry_paragraph(tbl._tbl.findall(NS + "tr")[i], frag)
        if p is None:
            continue
        seen_link = False
        for child in list(p):
            if child.tag == NS + "hyperlink":
                seen_link = True
            elif seen_link and child.tag == NS + "r" and text_of(child) in ("|", "preprint", ""):
                p.remove(child)
        report.append(f"  t{ti}: dropped the preprint label on {frag[:34]!r}")

    for ti, frag in DELETE:
        tbl = tables[ti]
        i = find_row(tbl, frag)
        trs = tbl._tbl.findall(NS + "tr")
        year_tc = copy.deepcopy(trs[i].findall(NS + "tc")[0])
        year = text_of(year_tc)
        trs[i].getparent().remove(trs[i])
        if year:                        # the year label must not leave with the row
            after = tbl._tbl.findall(NS + "tr")[i]
            if not text_of(after.findall(NS + "tc")[0]):
                after.replace(after.findall(NS + "tc")[0], year_tc)
                report.append(f"  t{ti}: moved the {year} label to the next entry")
        report.append(f"  t{ti}: removed the superseded preprint {frag[:40]!r}")

    for ti, new in ((0, SENIOR_NEW), (1, MIDDLE_NEW)):
        tbl = tables[ti]
        anchor = tbl._tbl.findall(NS + "tr")[0]
        template_row = copy.deepcopy(anchor)
        for year, authors, title, journal, url, suffix in new:
            tr = copy.deepcopy(template_row)
            tcs = tr.findall(NS + "tc")
            set_cell_text(tcs[0], year)
            set_cell_text(tcs[2], journal)
            build_entry(tcs[1].find(NS + "p"), tpl, authors, title, url, suffix, part)
            anchor.addprevious(tr)
            report.append(f"  t{ti}: added {title[:46]!r}")

    for ti, tbl in tables.items():                 # a year is labelled once per block
        seen = None
        for tr in tbl._tbl.findall(NS + "tr"):
            tcs = tr.findall(NS + "tc")
            if not tcs:
                continue
            y = text_of(tcs[0])
            if not y:
                continue
            if y == seen:
                for t in tcs[0].iter(NS + "t"):
                    t.text = ""
                report.append(f"  t{ti}: removed a repeated {y} label")
            else:
                seen = y

    # the header carries a last-saved-date field; set the property it reads so the
    # date is right once Word refreshes fields (Ctrl+A, F9, or on next save)
    d.core_properties.modified = dt.datetime(2026, 8, 11)
    d.core_properties.last_modified_by = "N Steinmetz"

    d.save(str(DST))
    print(f"wrote {DST}")
    print("\n".join(report))

    if "--pdf" in sys.argv:
        from docx2pdf import convert
        convert(str(DST), str(DST.with_suffix(".pdf")))
        print("rendered", DST.with_suffix(".pdf"))


if __name__ == "__main__":
    main()
